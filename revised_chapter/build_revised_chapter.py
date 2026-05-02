from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Chapter_1_Foundations_of_LLM_Applications_Revised.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D8DEE9", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(paragraph, text, bold=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_reference_paragraph(doc, text):
    paragraph = doc.add_paragraph(style="References")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.35)
    paragraph.add_run(text)


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    set_cell_border(cell, color="93C5FD")
    paragraph = cell.paragraphs[0]
    add_run(paragraph, title + " ", bold=True, color="1D4ED8")
    paragraph.add_run(body)
    doc.add_paragraph()


def add_flow_table(doc):
    table = doc.add_table(rows=2, cols=4)
    set_table_borders(table)
    headers = ["1. User Input", "2. Tokenization", "3. Model Processing", "4. Output"]
    details = [
        "The application receives a question, instruction, or document.",
        "Text is split into tokens that the model can process.",
        "The model uses learned patterns to estimate useful next tokens.",
        "Tokens are converted back into readable text for the user.",
    ]
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, "DBEAFE")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(paragraph, text, bold=True, color="1E3A8A")
    for index, text in enumerate(details):
        cell = table.cell(1, index)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
    doc.add_paragraph("Figure 1.1: A simplified LLM response flow.", style="Caption")


def add_comparison_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    header = table.rows[0].cells
    for i, text in enumerate(["Feature", "Traditional Software", "LLM-Based Application"]):
        set_cell_shading(header[i], "E0F2FE")
        p = header[i].paragraphs[0]
        add_run(p, text, bold=True, color="0F172A")
    rows = [
        ("Behavior", "Follows explicit rules written by developers.", "Generates responses using patterns learned from data."),
        ("Input handling", "Works best with predictable inputs.", "Handles many natural-language variations."),
        ("Output", "Usually deterministic for the same input.", "Can vary because generation is probabilistic."),
        ("Best use", "Structured workflows, calculations, and strict business logic.", "Language-heavy tasks such as summarization, drafting, and question answering."),
        ("Risk", "Limited flexibility.", "May produce inaccurate or uncited content if not reviewed."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
    doc.add_paragraph("Table 1.1: Comparing rule-based systems and LLM-based applications.", style="Caption")


def add_training_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    header = table.rows[0].cells
    for i, text in enumerate(["Aspect", "Training", "Inference"]):
        set_cell_shading(header[i], "F0FDF4")
        p = header[i].paragraphs[0]
        add_run(p, text, bold=True, color="14532D")
    rows = [
        ("Purpose", "Learn language patterns from large datasets.", "Use a trained model to respond to a user request."),
        ("When it happens", "Before deployment or during model improvement.", "During application use."),
        ("Compute needs", "Very high.", "Lower than training, but still can be significant."),
        ("Input", "Training data.", "User prompt or application-provided context."),
        ("Output", "Updated model parameters.", "Generated text, code, summary, or decision support."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
    doc.add_paragraph("Table 1.2: Training and inference in LLM systems.", style="Caption")


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    if "Caption" not in styles:
        styles.add_style("Caption", 1)
    styles["Caption"].font.name = "Arial"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor(75, 85, 99)
    if "References" not in styles:
        styles.add_style("References", 1)
    styles["References"].font.name = "Arial"
    styles["References"].font.size = Pt(9)

    header = section.header.paragraphs[0]
    header.text = "Foundations of LLM Applications"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    doc.add_paragraph("Chapter 1: Foundations of LLM Applications", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Revised chapter with original wording, cleaner formatting, and source references.").italic = True

    add_note_box(
        doc,
        "Revision note:",
        "This version fixes broken PDF characters, rewrites generic explanations, adds original project examples, and includes references for key LLM concepts.",
    )

    doc.add_heading("1.1 Introduction to LLM Applications", level=1)
    doc.add_paragraph(
        "Large Language Model (LLM) applications are software systems that connect a user interface, business logic, and a language model to support tasks involving human language. Instead of relying only on fixed rules, these applications can interpret prompts, use context, and generate useful text responses."
    )
    doc.add_paragraph(
        "A chatbot, writing assistant, coding helper, or document-review tool can all be designed as LLM applications. In each case, the application is responsible for collecting user input, preparing the request, sending it to a model, and presenting the model output in a useful way."
    )
    add_note_box(
        doc,
        "Project connection:",
        "In a plagiarism source finder, an LLM could explain why two sentences appear similar, while the search component provides source links and evidence for review.",
    )

    doc.add_heading("1.2 What Is a Large Language Model?", level=1)
    doc.add_paragraph(
        "A Large Language Model is an AI system trained on large collections of text so it can recognize language patterns and generate human-like responses. Stanford HAI describes LLMs as systems trained on massive text data to understand and generate language, while IBM explains that LLMs process prompts by converting text into smaller units and generating output one token at a time."
    )
    doc.add_paragraph(
        "The important idea is that an LLM does not understand text in the same way a person does. It processes numerical representations of tokens and predicts likely continuations based on patterns learned during training."
    )

    doc.add_heading("1.3 How LLM Applications Work", level=1)
    doc.add_paragraph(
        "Although real systems can include many components, the basic response flow is easy to understand: a user provides input, the input is prepared as tokens, the model processes those tokens, and the application returns output."
    )
    add_flow_table(doc)

    doc.add_heading("1.3.1 User Input", level=2)
    doc.add_paragraph(
        "The process begins when a user asks a question, uploads a document, or gives an instruction. The application may also add context, such as retrieved source material, user settings, or formatting requirements."
    )

    doc.add_heading("1.3.2 Tokenization", level=2)
    doc.add_paragraph(
        "Before text reaches the model, it is divided into tokens. A token can be a word, part of a word, punctuation mark, or other text unit. OpenAI and IBM both describe tokenization as the step that breaks text into units a model can process."
    )

    doc.add_heading("1.3.3 Model Processing", level=2)
    doc.add_paragraph(
        "After tokenization, the model processes the tokens using patterns stored in its parameters. The model estimates which tokens are likely to come next, guided by the input prompt and any context supplied by the application."
    )

    doc.add_heading("1.3.4 Output Generation", level=2)
    doc.add_paragraph(
        "The model generates output by producing tokens step by step. The application then converts those tokens back into readable text and may format, filter, or validate the response before showing it to the user."
    )

    doc.add_heading("1.4 LLMs vs. Traditional Software Systems", level=1)
    doc.add_paragraph(
        "Traditional software is usually built around explicit rules. If the same input enters the same deterministic function, the same output is expected. LLM-based applications are different because they generate language from learned patterns and can handle more flexible input."
    )
    add_comparison_table(doc)
    doc.add_paragraph(
        "This difference is powerful, but it also creates responsibility. Developers should design guardrails, review workflows, citations, and user-facing explanations so generated content can be trusted and evaluated."
    )

    doc.add_heading("1.5 Training vs. Inference", level=1)
    doc.add_paragraph(
        "LLM systems are often discussed in two phases: training and inference. Training is the learning phase, where a model is built or improved using large datasets. Inference is the usage phase, where the trained model responds to user input."
    )
    add_training_table(doc)
    doc.add_paragraph(
        "A simple analogy is studying and taking an exam. Training is similar to studying from many examples. Inference is similar to using that learned knowledge to answer a new question. The analogy is not perfect, but it helps separate model development from model usage."
    )

    doc.add_heading("1.6 Chapter Summary", level=1)
    doc.add_paragraph(
        "This chapter introduced the foundation of LLM applications. An LLM application receives user input, prepares it for a model, generates or processes language, and returns useful output. Unlike traditional software, LLM systems are flexible with language but require careful review, citations, and responsible design."
    )
    doc.add_paragraph(
        "For portfolio projects, the strongest LLM applications do more than call a model. They show clear user value, explain outputs, handle errors, cite sources when needed, and make the review process transparent."
    )

    doc.add_page_break()
    doc.add_heading("References", level=1)
    add_reference_paragraph(
        doc,
        "IBM. (2026). What are large language models (LLMs)? IBM Think. https://www.ibm.com/think/topics/large-language-models",
    )
    add_reference_paragraph(
        doc,
        "IBM. (2026). What is tokenization? IBM Think. https://www.ibm.com/think/topics/tokenization",
    )
    add_reference_paragraph(
        doc,
        "OpenAI. (2026). What are tokens and how to count them? OpenAI Help Center. https://help.openai.com/en/articles/4936856-what-are-tokens-and-how-to-count-them",
    )
    add_reference_paragraph(
        doc,
        "Stanford Institute for Human-Centered Artificial Intelligence. (2026). What is a large language model (LLM)? https://hai.stanford.edu/ai-definitions/what-is-a-llm",
    )
    add_reference_paragraph(
        doc,
        "Google for Developers. (2026). Machine Learning Glossary. https://developers.google.com/machine-learning/glossary",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
